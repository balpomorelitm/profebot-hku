import streamlit as st
import streamlit.components.v1 as components
import requests
import re
import time
import json
import os
import logging
import hashlib
from datetime import datetime, timedelta
from typing import Optional, List, Dict
from PIL import Image
from pathlib import Path
from collections import Counter
from io import BytesIO
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ==========================================
# LOGGING CONFIGURATION
# ==========================================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ==========================================
# CONFIGURATION & CONSTANTS
# ==========================================
# Paths
BASE_DIR = Path(__file__).parent
STYLES_DIR = BASE_DIR / "styles"
DATA_DIR = BASE_DIR / "data"
THREADS_FILE = DATA_DIR / "threads.json"
ANALYTICS_FILE = DATA_DIR / "analytics.json"
USER_PROFILE_FILE = DATA_DIR / "user_profile.json"
RESPONSE_CACHE_FILE = DATA_DIR / "response_cache.json"

# Cache configuration
CACHE_MAX_SIZE = 100
CACHE_TTL_HOURS = 168  # 1 week

# Ensure data directory exists
DATA_DIR.mkdir(exist_ok=True)

# Load favicon
try:
    favicon = Image.open("favicon.jpg")
except:
    favicon = "🎓"

st.set_page_config(
    page_title="ProfeBot - Spanish Year 1 Tutor",
    page_icon=favicon,
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# 🎨 UI CLEANUP (HIDE STREAMLIT CHROME)
# ==========================================
hide_streamlit_style = """
<style>
    /* Hide header elements but keep sidebar toggle visible */
    header[data-testid="stHeader"] {
        background-color: transparent !important;
    }
    
    header[data-testid="stHeader"] > div:first-child {
        background-color: transparent !important;
    }
    
    /* NUCLEAR OPTION: Hide everything in header except sidebar button */
    header[data-testid="stHeader"] > div:first-child > div {
        display: none !important;
    }
    
    /* Keep only the sidebar toggle button visible */
    header[data-testid="stHeader"] button[kind="header"] {
        display: block !important;
        visibility: visible !important;
    }
    
    /* Hide hamburger menu, deploy button, and other header items */
    #MainMenu {visibility: hidden;}
    .stDeployButton {display: none !important;}
    header .stActionButton {display: none !important;}
    
    /* Hide Fork button - comprehensive selectors */
    .stAppHeader a {display: none !important;}
    header a {display: none !important;}
    header iframe {display: none !important;}
    iframe[src*="github"] {display: none !important;}
    iframe[title] {display: none !important;}
    iframe {display: none !important;}
    a[href*="github.com"] {display: none !important;}
    a[href*="streamlit.io"] {display: none !important;}
    button[data-testid*="fork"] {display: none !important;}
    div[data-testid*="fork"] {display: none !important;}
    header > div > div > a {display: none !important;}
    [class*="viewerBadge"] {display: none !important;}
    [class*="GitHubBadge"] {display: none !important;}
    
    /* SIDEBAR STYLING */
    [data-testid="stSidebar"] {
        min-width: 18.47rem !important;
        max-width: 18.47rem !important;
        width: 18.47rem !important;
    }
    
    [data-testid="stSidebar"] > div:first-child {
        width: 18.47rem !important;
    }
    
    /* Make native collapse button always visible */
    [data-testid="collapsedControl"] {
        display: block !important;
        visibility: visible !important;
        opacity: 1 !important;
        position: fixed !important;
        left: 4px !important;
        top: 4px !important;
        z-index: 999999 !important;
        background: linear-gradient(135deg, #00A86B 0%, #0077C8 100%) !important;
        border: 2px solid rgba(255,255,255,0.3) !important;
        border-radius: 8px !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.3) !important;
        padding: 8px !important;
    }
    
    button[kind="header"] {
        display: block !important;
        visibility: visible !important;
        opacity: 1 !important;
    }
    
    /* Ensure button visible when collapsed */
    [data-testid="stSidebar"][aria-expanded="false"] ~ [data-testid="collapsedControl"] {
        display: block !important;
        visibility: visible !important;
        left: 4px !important;
    }
    
    /* Hide footer */
    footer {visibility: hidden;}
    
    /* Hide the 'Manage app' button and bottom toolbar */
    .stAppDeployButton {display: none !important;}
    [data-testid="stToolbar"] {visibility: hidden !important;}
    [data-testid="stStatusWidget"] {visibility: hidden !important;}
    [data-testid="stDecoration"] {visibility: hidden !important;}
    
    /* Hide GitHub link/button (bottom right) */
    .viewerBadge_container__1QSob {display: none !important;}
    .viewerBadge_link__1S137 {display: none !important;}
    .styles_viewerBadge__1yB5_ {display: none !important;}
    [data-testid="stAppViewBlockContainer"] > div:last-child {display: none !important;}
    iframe[title="streamlit_app"] {display: none !important;}
    
    /* SUPER AGGRESSIVE FORK BUTTON HIDING */
    header * {
        background-image: none !important;
    }
    
    header a, header iframe, header img {
        opacity: 0 !important;
        width: 0 !important;
        height: 0 !important;
        position: absolute !important;
        left: -9999px !important;
    }
    
    [data-testid="stAppViewBlockContainer"] a[href*="github"],
    [data-testid="stAppViewBlockContainer"] a[target="_blank"] {
        display: none !important;
        visibility: hidden !important;
        opacity: 0 !important;
    }
    
    /* Adjust top padding */
    .block-container {
        padding-top: 3rem !important; 
    }
    
    /* ===== MOBILE RESPONSIVE DESIGN ===== */
    @media (max-width: 768px) {
        /* HIDE sidebar completely on mobile */
        [data-testid="stSidebar"],
        [data-testid="collapsedControl"],
        button[kind="header"] {
            display: none !important;
        }
        
        /* Full width content on mobile */
        [data-testid="stAppViewContainer"] {
            margin-left: 0 !important;
        }
        
        .block-container {
            padding: 1rem !important;
            max-width: 100% !important;
        }
    }
    
    /* For very small screens (phones) */
    @media (max-width: 480px) {
        [data-testid="stSidebar"] {
            min-width: 14rem !important;
            max-width: 14rem !important;
            width: 14rem !important;
        }
        
        [data-testid="stSidebar"] > div:first-child {
            width: 14rem !important;
            padding: 0.5rem !important;
        }
        
        /* Smaller text in sidebar on very small screens */
        [data-testid="stSidebar"] {
            font-size: 0.85rem !important;
        }
        
        [data-testid="stSidebar"] h3 {
            font-size: 1rem !important;
        }
        
        [data-testid="stSidebar"] h4 {
            font-size: 0.9rem !important;
        }
    }
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# ==========================================
# JAVASCRIPT TO REMOVE FORK BUTTON
# ==========================================
import streamlit.components.v1 as components

remove_fork_script = """
<script>
// Remove Fork button with extreme prejudice
function removeForkButton() {
    // Remove all links in header
    const headerLinks = window.parent.document.querySelectorAll('header a');
    headerLinks.forEach(link => {
        link.style.display = 'none';
        link.remove();
    });
    
    // Remove all iframes
    const iframes = window.parent.document.querySelectorAll('header iframe');
    iframes.forEach(iframe => {
        iframe.style.display = 'none';
        iframe.remove();
    });
    
    // Remove any element containing "fork" or "github"
    const allElements = window.parent.document.querySelectorAll('header *');
    allElements.forEach(el => {
        const text = el.textContent?.toLowerCase() || '';
        const href = el.getAttribute('href')?.toLowerCase() || '';
        if (text.includes('fork') || text.includes('github') || href.includes('github')) {
            el.style.display = 'none';
            el.remove();
        }
    });
}

// Run immediately and on every rerun
removeForkButton();
setInterval(removeForkButton, 500);
</script>
"""

components.html(remove_fork_script, height=0)

# ==========================================
# CSS LOADING FROM FILES
# ==========================================
def load_css_from_file(dark_mode: bool = False) -> str:
    """Load CSS from external file based on theme."""
    css_file = STYLES_DIR / ("dark.css" if dark_mode else "light.css")
    try:
        with open(css_file, "r", encoding="utf-8") as f:
            css_content = f.read()
            logger.info(f"Loaded CSS from {css_file}")
            return f"<style>{css_content}</style>"
    except FileNotFoundError:
        logger.warning(f"CSS file not found: {css_file}, using fallback")
        return get_fallback_css(dark_mode)

def get_fallback_css(dark_mode: bool = False) -> str:
    """Minimal fallback CSS if files not found."""
    bg = "#0d1117" if dark_mode else "#ffffff"
    text = "#f0f6fc" if dark_mode else "#24292f"
    return f"""
    <style>
        .stApp {{ background-color: {bg}; color: {text}; }}
    </style>
    """

# API Configuration
MAX_RETRIES = 3
RETRY_DELAY = 2
REQUEST_TIMEOUT = 30

# Language configurations
LANGUAGE_OPTIONS = {
    "English": "English",
    "繁體中文 (Cantonese)": "Cantonese",
    "普通话 (Mandarin)": "Mandarin",
    "Other / Otro": "custom"
}

# ==========================================
# SECRET KEYS MANAGEMENT
# ==========================================
def load_secrets() -> Dict[str, str]:
    """Load secrets with fallback for local testing."""
    try:
        hku_api_key = st.secrets["HKU_API_KEY"]
        # Try to get GPT key, fallback to HKU_API_KEY if not found
        try:
            hku_gpt_key = st.secrets["HKU_GPT_KEY"]
        except KeyError:
            hku_gpt_key = hku_api_key
            logger.warning("HKU_GPT_KEY not found, using HKU_API_KEY as fallback")
        
        return {
            "NOTION_TOKEN": st.secrets["NOTION_TOKEN"],
            "DATABASE_ID": st.secrets["DATABASE_ID"],
            "HKU_API_KEY": hku_api_key,
            "HKU_GPT_KEY": hku_gpt_key
        }
    except (FileNotFoundError, KeyError) as e:
        st.sidebar.error("⚠️ Missing secrets configuration")
        return {
            "NOTION_TOKEN": "your_notion_token_here",
            "DATABASE_ID": "your_database_id_here",
            "HKU_API_KEY": "your_hku_api_key_here",
            "HKU_GPT_KEY": "your_hku_gpt_key_here"
        }

secrets = load_secrets()
NOTION_TOKEN = secrets["NOTION_TOKEN"]
DATABASE_ID = secrets["DATABASE_ID"]

# ==========================================
# HYBRID ROUTER - MODEL CONFIGURATION
# ==========================================
# API Base URL for HKU
HKU_API_BASE = "https://api.hku.hk"
HKU_API_VERSION = "2025-01-01-preview"

# Fast Model (DeepSeek-V3) - For routing and simple queries
MODEL_FAST_ID = "DeepSeek-V3"
ENDPOINT_FAST = f"{HKU_API_BASE}/deepseek/models/chat/completions"
KEY_FAST = secrets["HKU_API_KEY"]

# Smart Model (GPT-4.1) - For complex queries
# Available models: gpt-4.1-nano, gpt-4.1-mini, gpt-4.1, o4-mini, gpt-5-nano, gpt-5-mini, gpt-5-chat, gpt-5, gpt-5.1
MODEL_SMART_ID = "gpt-4.1"  # Good balance of quality and speed
ENDPOINT_SMART = f"{HKU_API_BASE}/openai/deployments/{MODEL_SMART_ID}/chat/completions?api-version={HKU_API_VERSION}"
KEY_SMART = secrets["HKU_API_KEY"]  # Same key for all HKU APIs

# Enable hybrid routing (uses GPT for complex queries, DeepSeek for simple)
USE_HYBRID_ROUTER = True

# Legacy compatibility
DEPLOYMENT_ID = MODEL_FAST_ID
HKU_API_KEY = KEY_FAST
HKU_ENDPOINT = ENDPOINT_FAST

# ==========================================
# HELPER FUNCTIONS
# ==========================================
def make_request_with_retry(
    method: str,
    url: str,
    headers: dict,
    json_payload: Optional[dict] = None,
    params: Optional[dict] = None,
    max_retries: int = MAX_RETRIES
) -> Optional[requests.Response]:
    """Make HTTP request with automatic retry logic for transient errors."""
    for attempt in range(max_retries):
        try:
            if method.upper() == "POST":
                response = requests.post(
                    url, 
                    headers=headers, 
                    json=json_payload, 
                    params=params,
                    timeout=REQUEST_TIMEOUT
                )
            elif method.upper() == "GET":
                response = requests.get(
                    url, 
                    headers=headers, 
                    params=params,
                    timeout=REQUEST_TIMEOUT
                )
            else:
                raise ValueError(f"Unsupported HTTP method: {method}")
            
            if response.status_code in [502, 503, 504]:
                if attempt < max_retries - 1:
                    time.sleep(RETRY_DELAY * (attempt + 1))
                    continue
            
            return response
            
        except requests.exceptions.Timeout:
            if attempt < max_retries - 1:
                time.sleep(RETRY_DELAY * (attempt + 1))
                continue
            else:
                st.error(f"⏱️ Request timeout after {max_retries} attempts")
                return None
                
        except requests.exceptions.ConnectionError as e:
            if attempt < max_retries - 1:
                time.sleep(RETRY_DELAY * (attempt + 1))
                continue
            else:
                st.error(f"🔌 Connection error after {max_retries} attempts")
                logger.error(f"Connection error to {url}: {str(e)}")
                return None
                
        except Exception as e:
            st.error(f"❌ Unexpected error: {str(e)}")
            return None
    
    return None

def generate_thread_title(first_message: str) -> str:
    """Generate a short title from the first user message."""
    clean_msg = first_message.strip()
    if len(clean_msg) > 30:
        return clean_msg[:30] + "..."
    return clean_msg

def get_language_instruction(language: str, custom_language: str = "") -> str:
    """Get language-specific instruction for the prompt."""
    if language == "English":
        return "All grammatical explanations, feedback, and answers to administrative questions MUST be in ENGLISH."
    elif language == "Cantonese":
        return "All grammatical explanations, feedback, and answers to administrative questions MUST be in CANTONESE (繁體中文 - 粵語)."
    elif language == "Mandarin":
        return "All grammatical explanations, feedback, and answers to administrative questions MUST be in MANDARIN (普通话 - 简体中文)."
    elif language == "custom" and custom_language:
        return f"All grammatical explanations, feedback, and answers to administrative questions MUST be in {custom_language.upper()}."
    else:
        return "All grammatical explanations, feedback, and answers to administrative questions MUST be in ENGLISH."

def get_current_semester_info() -> dict:
    """Determine current semester based on month and return relevant unit information."""
    current_month = datetime.now().month
    
    if 1 <= current_month <= 5:  # January to May = Semester 2
        return {
            "semester": 2,
            "semester_name": "Second Semester (Spring)",
            "primary_units": "Units 7-12",
            "secondary_units": "Units 1-4 (next level book)",
            "focus_description": "The student is in the SECOND SEMESTER. Prioritize content from Units 7 onwards. Units 0-6 were covered last semester and can be referenced for review, but new content should focus on Units 7+.",
            "months": "January - May"
        }
    elif 9 <= current_month <= 12:  # September to December = Semester 1
        return {
            "semester": 1,
            "semester_name": "First Semester (Fall)",
            "primary_units": "Units 0-6",
            "secondary_units": "None (first semester)",
            "focus_description": "The student is in the FIRST SEMESTER. Focus primarily on Units 0-6. Do not introduce content from Units 7+ as they haven't covered that material yet.",
            "months": "September - December"
        }
    else:  # June to August = Summer break / transition period
        return {
            "semester": 0,
            "semester_name": "Summer Period",
            "primary_units": "All units for review",
            "secondary_units": "Preparation for next semester",
            "focus_description": "This is the SUMMER PERIOD between semesters. The student may be reviewing material or preparing for the next academic year. Cover all units as needed based on the student's questions.",
            "months": "June - August"
        }

# ==========================================
# USER PROFILE & MEMORY SYSTEM
# ==========================================
def load_user_profile() -> Dict:
    """Load user learning profile for personalization."""
    try:
        if USER_PROFILE_FILE.exists():
            with open(USER_PROFILE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as e:
        logger.error(f"Error loading user profile: {e}")
    return {
        "name": None,
        "weak_areas": [],
        "strong_areas": [],
        "vocabulary_errors": {},
        "grammar_errors": {},
        "completed_units": [],
        "quiz_scores": [],
        "learning_streak": 0,
        "last_active": None,
        "total_interactions": 0,
        "favorite_topics": {},
        "achievements": []
    }

def save_user_profile(profile: Dict):
    """Save user learning profile."""
    try:
        profile["last_active"] = datetime.now().isoformat()
        with open(USER_PROFILE_FILE, "w", encoding="utf-8") as f:
            json.dump(profile, f, ensure_ascii=False, indent=2)
        logger.info("User profile saved")
    except Exception as e:
        logger.error(f"Error saving user profile: {e}")

def update_learning_streak(profile: Dict) -> Dict:
    """Update the user's learning streak based on activity."""
    today = datetime.now().date()
    last_active = profile.get("last_active")
    
    if last_active:
        try:
            last_date = datetime.fromisoformat(last_active).date()
            days_diff = (today - last_date).days
            
            if days_diff == 0:
                pass  # Same day, streak unchanged
            elif days_diff == 1:
                profile["learning_streak"] = profile.get("learning_streak", 0) + 1
            else:
                profile["learning_streak"] = 1  # Reset streak
        except:
            profile["learning_streak"] = 1
    else:
        profile["learning_streak"] = 1
    
    return profile

def track_user_interaction(user_message: str, ai_response: str):
    """Analyze interaction to track learning patterns and update profile."""
    profile = load_user_profile()
    
    # Update streak
    profile = update_learning_streak(profile)
    
    # Increment total interactions
    profile["total_interactions"] = profile.get("total_interactions", 0) + 1
    
    # Detect correction patterns in AI response (indicates areas to improve)
    correction_indicators = [
        (r"(?:actually|correction|careful|remember|note that|be careful)", "general"),
        (r"(?:común error|common mistake|incorrecto|incorrect)", "general"),
        (r"(?:ser|estar)", "ser_estar"),
        (r"(?:género|gender|masculin|feminin)", "gender"),
        (r"(?:conjugat|conjugación)", "conjugation"),
        (r"(?:artículo|article|el |la |los |las )", "articles"),
        (r"(?:preposición|preposition|por|para)", "prepositions"),
        (r"(?:accent|acento|tilde)", "accents"),
    ]
    
    response_lower = ai_response.lower()
    message_lower = user_message.lower()
    
    # Check for corrections/errors
    has_correction = bool(re.search(r"(?:actually|correction|careful|incorrecto|mistake)", response_lower))
    
    if has_correction:
        for pattern, topic in correction_indicators:
            if re.search(pattern, response_lower) or re.search(pattern, message_lower):
                if topic != "general":
                    profile["grammar_errors"][topic] = profile["grammar_errors"].get(topic, 0) + 1
    
    # Track favorite topics based on questions
    topic_keywords = {
        "vocabulary": ["vocabulario", "vocabulary", "word", "palabra", "meaning", "significa"],
        "grammar": ["gramática", "grammar", "rule", "regla"],
        "conjugation": ["conjugat", "verb", "verbo"],
        "pronunciation": ["pronuncia", "sound", "accent", "sonido"],
        "culture": ["cultura", "culture", "spain", "españa", "mexico", "país"],
        "exercises": ["ejercicio", "exercise", "practice", "quiz", "task", "práctica"]
    }
    
    for topic, keywords in topic_keywords.items():
        if any(kw in message_lower for kw in keywords):
            profile["favorite_topics"][topic] = profile["favorite_topics"].get(topic, 0) + 1
    
    # Track quiz scores if detected
    score_patterns = [
        r"(\d+)\s*/\s*(\d+)",
        r"(\d+)\s*(?:out of|de)\s*(\d+)",
        r"(?:score|puntuación):\s*(\d+)\s*/\s*(\d+)",
    ]
    
    for pattern in score_patterns:
        match = re.search(pattern, response_lower)
        if match:
            try:
                correct = int(match.group(1))
                total = int(match.group(2))
                if total > 0 and correct <= total:
                    profile["quiz_scores"].append({
                        "correct": correct,
                        "total": total,
                        "percentage": round(correct / total * 100, 1),
                        "date": datetime.now().isoformat()
                    })
                    profile["quiz_scores"] = profile["quiz_scores"][-50:]  # Keep last 50
                    break
            except:
                pass
    
    save_user_profile(profile)
    return profile

def get_user_context_for_prompt() -> str:
    """Generate a context string about the user for the AI prompt."""
    profile = load_user_profile()
    
    context_parts = []
    
    # Learning streak
    streak = profile.get("learning_streak", 0)
    if streak > 0:
        context_parts.append(f"The student has a {streak}-day learning streak.")
    
    # Weak areas
    grammar_errors = profile.get("grammar_errors", {})
    if grammar_errors:
        sorted_errors = sorted(grammar_errors.items(), key=lambda x: x[1], reverse=True)[:3]
        weak_topics = [topic for topic, _ in sorted_errors]
        if weak_topics:
            context_parts.append(f"Areas needing extra attention: {', '.join(weak_topics)}.")
    
    # Quiz performance
    quiz_scores = profile.get("quiz_scores", [])
    if quiz_scores:
        recent_scores = quiz_scores[-5:]
        avg_score = sum(s["percentage"] for s in recent_scores) / len(recent_scores)
        if avg_score < 60:
            context_parts.append(f"Recent quiz average: {avg_score:.0f}% - student may need more practice.")
        elif avg_score >= 90:
            context_parts.append(f"Recent quiz average: {avg_score:.0f}% - student is performing excellently!")
    
    # Favorite topics
    fav_topics = profile.get("favorite_topics", {})
    if fav_topics:
        top_topic = max(fav_topics.items(), key=lambda x: x[1])[0]
        context_parts.append(f"Student shows high interest in: {top_topic}.")
    
    if context_parts:
        return "\n[USER LEARNING PROFILE]\n" + "\n".join(context_parts)
    return ""

# ==========================================
# RESPONSE CACHE SYSTEM
# ==========================================
def generate_cache_key(question: str, language: str) -> str:
    """Generate a hash key for caching purposes."""
    # Normalize the question
    normalized = question.lower().strip()
    normalized = re.sub(r'\s+', ' ', normalized)
    normalized = re.sub(r'[?!.,;:]', '', normalized)
    # Remove CMD_ prefixes for better matching
    normalized = re.sub(r'cmd_\w+:\s*', '', normalized)
    cache_key = f"{normalized}|{language}"
    return hashlib.md5(cache_key.encode()).hexdigest()

def get_cached_response(question: str, language: str) -> Optional[str]:
    """Get cached response if available and not expired."""
    try:
        if not RESPONSE_CACHE_FILE.exists():
            return None
        
        question_hash = generate_cache_key(question, language)
        
        with open(RESPONSE_CACHE_FILE, "r", encoding="utf-8") as f:
            cache = json.load(f)
        
        if question_hash in cache:
            entry = cache[question_hash]
            cached_time = datetime.fromisoformat(entry["timestamp"])
            
            if datetime.now() - cached_time < timedelta(hours=CACHE_TTL_HOURS):
                logger.info(f"Cache HIT for question: {question[:50]}...")
                return entry["response"]
            else:
                logger.info(f"Cache EXPIRED for question: {question[:50]}...")
    except Exception as e:
        logger.error(f"Cache read error: {e}")
    
    return None

def cache_response(question: str, language: str, response: str):
    """Cache a response for future use."""
    # Don't cache error responses or very short responses
    if "❌" in response or len(response) < 100:
        return
    
    # Don't cache dynamic content (quizzes with random elements, etc.)
    dynamic_indicators = ["CMD_QUIZ", "CMD_TASKS", "CMD_ROLEPLAY"]
    if any(indicator in question.upper() for indicator in dynamic_indicators):
        return
    
    try:
        cache = {}
        if RESPONSE_CACHE_FILE.exists():
            with open(RESPONSE_CACHE_FILE, "r", encoding="utf-8") as f:
                cache = json.load(f)
        
        question_hash = generate_cache_key(question, language)
        
        cache[question_hash] = {
            "question_preview": question[:200],
            "response": response,
            "language": language,
            "timestamp": datetime.now().isoformat(),
            "hit_count": cache.get(question_hash, {}).get("hit_count", 0) + 1
        }
        
        # Prune old entries if cache is too large
        if len(cache) > CACHE_MAX_SIZE:
            # Sort by timestamp and keep most recent
            sorted_entries = sorted(
                cache.items(), 
                key=lambda x: x[1].get("timestamp", ""),
                reverse=True
            )
            cache = dict(sorted_entries[:CACHE_MAX_SIZE])
        
        with open(RESPONSE_CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(cache, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Cached response for: {question[:50]}...")
            
    except Exception as e:
        logger.error(f"Cache write error: {e}")

def get_cache_stats() -> Dict:
    """Get cache statistics for display."""
    try:
        if RESPONSE_CACHE_FILE.exists():
            with open(RESPONSE_CACHE_FILE, "r", encoding="utf-8") as f:
                cache = json.load(f)
            
            total_entries = len(cache)
            total_hits = sum(entry.get("hit_count", 0) for entry in cache.values())
            
            return {
                "entries": total_entries,
                "total_hits": total_hits,
                "max_size": CACHE_MAX_SIZE
            }
    except:
        pass
    return {"entries": 0, "total_hits": 0, "max_size": CACHE_MAX_SIZE}

# ==========================================
# NOTION CONNECTION WITH CACHING
# ==========================================
@st.cache_data(ttl=3600)
def get_weekly_content() -> str:
    """Fetch active content from Notion database with caching."""
    url = f"https://api.notion.com/v1/databases/{DATABASE_ID}/query"
    headers = {
        "Authorization": f"Bearer {NOTION_TOKEN}",
        "Notion-Version": "2022-06-28",
        "Content-Type": "application/json"
    }
    payload = {
        "filter": {"property": "Activo", "checkbox": {"equals": True}}
    }

    response = make_request_with_retry("POST", url, headers, json_payload=payload)
    
    if not response:
        return "❌ Failed to connect to Notion after multiple attempts."
    
    if response.status_code != 200:
        return f"❌ Notion API Error ({response.status_code}): {response.text[:200]}"
    
    try:
        data = response.json()
        results = data.get("results", [])
        
        if not results:
            logger.warning("No active units found in Notion database")
            return "⚠️ No active units found in database."

        full_context = ""
        for page in results:
            props = page.get("properties", {})
            
            def get_text_safe(col_name: str) -> str:
                """Safely extract text from Notion properties with validation."""
                try:
                    if col_name == "Nombre":
                        prop_data = props.get("Nombre", {})
                        items = prop_data.get("title", []) if prop_data else []
                    else:
                        prop_data = props.get(col_name, {})
                        items = prop_data.get("rich_text", []) if prop_data else []
                    
                    if not items:
                        logger.debug(f"No content found for column: {col_name}")
                        return ""
                    
                    texts = []
                    for item in items:
                        if isinstance(item, dict):
                            text_content = item.get("text", {})
                            if isinstance(text_content, dict):
                                content = text_content.get("content", "")
                                if content:
                                    texts.append(content)
                    return " ".join(texts)
                except Exception as e:
                    logger.error(f"Error extracting text from {col_name}: {e}")
                    return ""

            name = get_text_safe("Nombre")
            lexicon = get_text_safe("Léxico")
            grammar = get_text_safe("Gramática")
            tags = get_text_safe("Tags")
            exercises = get_text_safe("Ejercicios")

            if name:  # Only add unit if it has a name
                full_context += f"""
=== UNIT: {name} ===
[TAGS]: {tags or 'No tags listed'}
[VOCABULARY]: {lexicon or 'No vocabulary listed'}
[GRAMMAR]: {grammar or 'No grammar listed'}
[APPROVED EXERCISES]: {exercises or 'No exercises listed'}
==============================
"""
                logger.info(f"Loaded unit: {name}")
        
        if not full_context:
            return "⚠️ No valid units found in database."
            
        return full_context
        
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error: {e}")
        return f"❌ Error parsing Notion response: Invalid JSON"
    except Exception as e:
        logger.error(f"Error parsing Notion data: {e}")
        return f"❌ Error parsing Notion data: {str(e)}"

# ==========================================
# AI CONNECTION - HYBRID ROUTER SYSTEM
# ==========================================
def call_ai_model(
    messages: List[Dict],
    model_type: str = "fast",
    max_tokens: int = 1000,
    temperature: float = 0.4
) -> Optional[str]:
    """Call AI model with appropriate configuration for DeepSeek or Azure OpenAI.
    
    Args:
        messages: List of message dicts with 'role' and 'content'
        model_type: 'fast' for DeepSeek-V3, 'smart' for GPT-5.1
        max_tokens: Maximum tokens in response
        temperature: Response creativity (0-1)
    
    Returns:
        Response content string or None if failed
    """
    if model_type == "smart":
        # HKU OpenAI (GPT-4.1) configuration
        # Uses same Ocp-Apim-Subscription-Key header as DeepSeek
        headers = {
            "Content-Type": "application/json",
            "Cache-Control": "no-cache",
            "Ocp-Apim-Subscription-Key": KEY_SMART
        }
        payload = {
            "messages": messages,
            "max_tokens": max_tokens,
            "temperature": temperature
        }
        response = make_request_with_retry(
            "POST",
            ENDPOINT_SMART,
            headers,
            json_payload=payload
        )
    else:
        # DeepSeek-V3 configuration
        headers = {
            "Content-Type": "application/json",
            "Cache-Control": "no-cache",
            "Ocp-Apim-Subscription-Key": KEY_FAST
        }
        payload = {
            "model": MODEL_FAST_ID,
            "messages": messages,
            "max_tokens": max_tokens,
            "temperature": temperature
        }
        response = make_request_with_retry(
            "POST",
            ENDPOINT_FAST,
            headers,
            json_payload=payload,
            params={"deployment-id": MODEL_FAST_ID}
        )
    
    if not response:
        logger.error(f"No response from {model_type} model")
        return None
    
    if response.status_code == 200:
        try:
            return response.json()['choices'][0]['message']['content']
        except (KeyError, IndexError) as e:
            logger.error(f"Unexpected API response format from {model_type}: {e}")
            return None
    else:
        logger.error(f"{model_type} API Error ({response.status_code}): {response.text[:200]}")
        return None


def classify_query_complexity(user_message: str, conversation_history: List[Dict] = None) -> tuple:
    """Use fast model to classify query as SIMPLE or COMPLEX.
    
    Returns:
        Tuple of (classification: str, reasoning: str)
    """
    # Build context from recent history
    history_context = ""
    if conversation_history:
        recent_msgs = conversation_history[-4:]  # Last 4 messages for context
        history_context = "\n".join([f"{m['role']}: {m['content'][:200]}" for m in recent_msgs])
    
    router_prompt = f"""You are a query classifier for a Spanish language tutor chatbot. Analyze the user's query and classify it.

CLASSIFY AS "SIMPLE" if the query is:
- Basic vocabulary questions ("What does X mean?")
- Simple greetings or casual conversation
- Requests for examples of words they already know
- Yes/no questions about Spanish basics
- Requests to repeat or clarify previous information
- Administrative questions about the course

CLASSIFY AS "COMPLEX" if the query is:
- Grammar explanations requiring detailed breakdowns
- Requests for tasks, quizzes, or exercises (CMD_TASKS, CMD_QUIZ)
- Comparisons between grammatical structures
- Cultural explanations requiring nuance
- Reading comprehension tasks
- Roleplay or conversation practice (CMD_ROLEPLAY)
- Questions about verb conjugation patterns
- Requests for "more explanation" (CMD_EXPLAIN_MORE)
- Any multi-step or creative content generation

Recent conversation context:
{history_context}

User query: "{user_message}"

Respond with ONLY one word: SIMPLE or COMPLEX"""

    messages = [{"role": "user", "content": router_prompt}]
    
    result = call_ai_model(messages, model_type="fast", max_tokens=10, temperature=0.1)
    
    if result:
        classification = "COMPLEX" if "COMPLEX" in result.upper() else "SIMPLE"
        return classification
    
    # Default to SIMPLE if router fails
    logger.warning("Router classification failed, defaulting to SIMPLE")
    return "SIMPLE"


def get_ai_response(user_message: str, notion_context: str, language: str, custom_language: str = "", conversation_history: List[Dict] = None) -> str:
    """Get AI response from HKU API with error handling and conversation history.
    
    Args:
        user_message: The current user message
        notion_context: The course content from Notion
        language: Preferred language for explanations
        custom_language: Custom language if 'Other' selected
        conversation_history: List of previous messages in the conversation
    """
    
    # Check cache first for simple, non-contextual queries
    is_contextual = conversation_history and len(conversation_history) > 2
    if not is_contextual:
        cached = get_cached_response(user_message, language)
        if cached:
            # Add cache indicator for router info
            return f"{cached}\n<!--ROUTER_DEBUG:CACHED|Cache-->"
    
    language_instruction = get_language_instruction(language, custom_language)
    semester_info = get_current_semester_info()
    user_context = get_user_context_for_prompt()
    
    system_prompt = f"""
[ROLE AND PROFILE]
You are "ProfeBot", the official Spanish Tutor for Spanish Year 1 at the University of Hong Kong (HKU).

**WHO YOU ARE:**
- An exceptional Spanish as a Foreign Language (ELE) teacher with a passion for making Spanish accessible and fun
- Expert at breaking down complex grammar into simple, digestible explanations
- Creative with examples: you use real-life situations, humor, and memorable scenarios to help concepts stick
- Culturally aware: you sprinkle in interesting facts about Spanish-speaking countries when relevant
- Patient and encouraging: you celebrate small wins and never make students feel bad for mistakes

**YOUR TEACHING STYLE:**
- You give CLEAR, PRACTICAL examples that students can immediately use in real conversations
- You anticipate common mistakes and address them proactively
- You use mnemonics, patterns, and comparisons to help students remember
- You connect new concepts to what students already know
- You make grammar feel logical, not arbitrary
- When explaining, you don't just give rules - you explain the WHY behind them

**YOUR STUDENTS:**
- Adults at HKU who are intelligent and multilingual (English, Mandarin, Cantonese)
- Complete beginners in Spanish (A1 level)
- Busy university students who appreciate efficient, focused explanations

**YOUR TONE:**
- Academic yet approachable
- Encouraging and positive (use emojis sparingly but warmly 😊)
- Patient with repetition - if they ask the same thing twice, explain it a different way
- Enthusiastic about Spanish without being overwhelming

[⚠️ CRITICAL LANGUAGE PROTOCOL - MANDATORY ⚠️]
The student's preferred language setting is: **{language_instruction}**

**ABSOLUTE RULES - NEVER VIOLATE:**
1. ALL your text MUST be written in the STUDENT'S PREFERRED LANGUAGE, including:
   - Explanations and feedback
   - Instructions for tasks, quizzes, and exercises
   - Questions you ask the student
   - Multiple choice options text (except Spanish vocabulary being tested)
   - The 3 follow-up suggestions at the end (///)
   - ANY administrative or conversational text

2. **ONLY USE SPANISH FOR:**
   - Spanish vocabulary words/phrases being taught or tested
   - Example sentences demonstrating Spanish usage
   - Reading passages in reading tasks (the text itself, NOT the questions)
   - The Spanish words in fill-in-the-blank exercises

3. **NEVER write instructions, questions, feedback, or explanations in Spanish** unless the student explicitly set Spanish as their preferred language.

4. If a student writes in Chinese, understand it but ALWAYS respond in their PREFERRED LANGUAGE.

[CONTENT SACRED RULES]
1. THE "ACTIVE CONTENT" IS YOUR BIBLE:
   - You are STRICTLY FORBIDDEN from using vocabulary, verb tenses, or grammar rules that do not appear in the "ACTIVE CONTENT" list below.
   - If the student asks about something advanced (e.g., "fui" - past tense), congratulate their curiosity but explain IN THEIR PREFERRED LANGUAGE that it belongs to future levels and rephrase using ONLY what they know now.

[🎯 SMART ROUTING SYSTEM - USE TAGS]
When a student asks a question, follow this process:
1. **FIRST**: Read the [TAGS] section of each unit - these contain key concepts, themes, and topics covered in that unit.
2. **IDENTIFY**: Match the student's question to the most relevant unit(s) based on the tags.
3. **THEN**: Look into the [VOCABULARY] and [GRAMMAR] sections of the identified unit(s) to provide accurate, curriculum-aligned answers.
4. **EXAMPLE**: If a student asks about "numbers" or "age", check which unit has tags related to numbers/age, then use ONLY the vocabulary and grammar from that unit.
5. **MULTIPLE UNITS**: If the topic spans multiple units, combine information from all relevant units but clearly indicate which content comes from which unit.

[📅 SEMESTER-AWARE CONTENT PRIORITIZATION]
**Current Academic Period:** {semester_info['semester_name']} ({semester_info['months']})
**Primary Focus Units:** {semester_info['primary_units']}
**Secondary/Advanced Units:** {semester_info['secondary_units']}

**IMPORTANT CONTEXT:** {semester_info['focus_description']}

**How to apply this:**
- When answering questions, PRIORITIZE content from the primary focus units for this semester.
- If a student asks about content from units they haven't covered yet (based on the semester), gently explain that this topic will be covered later in the course, but you can give them a brief preview if they're curious.
- For review questions about past semesters' content, feel free to help but remind them that this was previous material.
- Always check which units are currently "Active" in the database - these represent what the teacher has enabled for the current period.
{user_context}

[TASK GENERATION SYSTEM]
When the user requests a TASK (CMD_TASKS), follow this protocol:
⚠️ LANGUAGE REMINDER: ALL instructions, questions, and feedback MUST be in the STUDENT'S PREFERRED LANGUAGE. Only the Spanish text/vocabulary being practiced should be in Spanish.

1. **First, ASK the student** (IN THEIR PREFERRED LANGUAGE) which type of task they want:
   - **Reading Task**: Create a 200-word text IN SPANISH. Then provide 8 multiple choice comprehension questions WITH ALL INSTRUCTIONS AND OPTIONS IN THE STUDENT'S PREFERRED LANGUAGE.
   - **Conversation Task**: Generate conversation prompts. ALL INSTRUCTIONS must be in the student's preferred language.
   - **Grammar & Vocabulary Task**: Design an exercise. ALL INSTRUCTIONS AND FEEDBACK must be in the student's preferred language.
2. **Ask which unit** they want to practice (in their preferred language).
3. **Wait for their response** before creating the task.

[QUIZ GENERATION LOGIC]
When the user requests a QUIZ (CMD_QUIZ):
⚠️ LANGUAGE REMINDER: Write ALL quiz instructions, questions, and feedback in the STUDENT'S PREFERRED LANGUAGE.

1. **First, ASK the student** (IN THEIR PREFERRED LANGUAGE) what topic or vocabulary they want to practice.
2. **IMPORTANT**: When you provide quiz questions, do NOT include the answers.
3. ALL question prompts and instructions = STUDENT'S PREFERRED LANGUAGE. Only Spanish vocabulary being tested = Spanish.
4. Wait for the student to submit their answers.
5. Provide detailed feedback IN THE STUDENT'S PREFERRED LANGUAGE.

**⚠️ MANDATORY QUIZ FORMAT** (for interactive quiz system):
When generating multiple choice questions, you MUST use this EXACT format:

1. Question text here?
A) First option
B) Second option
C) Third option
D) Fourth option (optional)

2. Next question here?
A) Option A
B) Option B
C) Option C

Rules for quiz format:
- Each question MUST start with a number followed by a period or parenthesis (e.g., "1." or "1)")
- Each option MUST be on its own line
- Options MUST use capital letters A, B, C, D followed by a closing parenthesis (e.g., "A)")
- Leave a blank line between questions
- Do NOT use bold/markdown in the question numbers or option letters

[EXERCISE GENERATION LOGIC]
⚠️ LANGUAGE REMINDER: ALL exercise instructions and explanations MUST be in the STUDENT'S PREFERRED LANGUAGE.

- When the user asks for an exercise or practice:
  1. **PRIORITY 1**: Check the `[APPROVED EXERCISES]` section in the content below. If there are exercises there, USE THEM.
  2. **PRIORITY 2**: If NO exercises are listed there (or they are exhausted), generate a meaningful "Fill-in-the-blanks" or "Multiple Choice" exercise based STRICTLY on the vocabulary list.
  3. **QUALITY CONTROL**: AVOID trivial questions (e.g., "Is 'Hola' written with H?"). Create communicative context (e.g., "Complete the dialogue between A and B").

[EXPLAIN MORE COMMAND]
When the user requests CMD_EXPLAIN_MORE:
- Elaborate on the topic you were just discussing IN THE STUDENT'S PREFERRED LANGUAGE.
- Go slightly deeper, provide additional context, examples, or nuances.
- Keep explanations appropriate to the student's level (beginner).
- Do not introduce advanced grammar or vocabulary outside the active content.

[EXTERNAL LEARNING TOOLS & RESOURCES]
⚠️ LANGUAGE REMINDER: Explain all resources IN THE STUDENT'S PREFERRED LANGUAGE.

**OFFICIAL COURSE RESOURCES - Recommend these when relevant:**
When students ask about:
- **Course syllabus, schedule, or course structure** → Recommend: [📋 Course Syllabus 2025/26](https://hkuhk-my.sharepoint.com/:w:/g/personal/pablot_hku_hk/IQA9JsO9FSBJQLiOPSfk4w-lAWRCn6skb-ObwSR_vtj4cZk?e=ULXiqR)

- **Course content overview, what topics are covered, or semester planning** → Recommend: [📖 Course Contents Summary](https://hkuhk-my.sharepoint.com/:w:/g/personal/pablot_hku_hk/IQDINc5UzBQhS4Mp4t2rC99_AWaikqT-OGpC4-9vaKgY7wM?e=uJyeHa) - Comprehensive English summary of all course content

- **Assignments, grades, course materials, or general course questions** → Recommend: [🏠 Moodle Course 2025/26](https://moodle.hku.hk/course/view.php?id=136141)

- **Latest updates, announcements, important dates** → Recommend: [📢 Course Announcements](https://moodle.hku.hk/mod/forum/view.php?id=3990047)

- **Extra reading materials, practice books, or wanting to read more in Spanish** → Recommend: [📕 Easy Readers (e-books)](https://moodle.hku.hk/pluginfile.php/6225750/mod_folder/content/0/Easy%20readers%20%28e-books%29.pdf?forcedownload=1) - Great for additional practice

**How to recommend resources:**
- Mention the resource naturally in your response (in the student's preferred language)
- Include the clickable markdown link
- Briefly explain what they'll find there
- Example: "Para ver el cronograma completo del curso, puedes revisar el [📋 Course Syllabus 2025/26](link). Allí encontrarás todas las fechas importantes."

When students ask about external tools, apps, games, or resources to practice Spanish:
1. **FIRST PRIORITY - HKU TEACHER-DESIGNED GAMES**: Enthusiastically recommend the digital games created specifically for Spanish Year 1 by your teachers:
   - 🤖 **The CONJUGATOR**: A game designed by HKU teachers to practice verb conjugation: https://conjugator.pablotorrado.site/
   - 🟧🟩⬜ **Palabrero**: The daily Wordle of the course - practice vocabulary every day! https://span1001palabrero.netlify.app/
   - 🚢🧨 **Batalla Verbal (Battleship)**: A conjugation pairs game based on Battleship https://batallaverbal.netlify.app/
   - 🗺️ **Mapamundi**: A world geography game based on Unit 3 contents https://balpomorelitm.github.io/mapamundi/
   - All these games are available in the "General Information" section of the course Notion page

2. **ADDITIONAL RECOMMENDED RESOURCES** (mention these when relevant):
   - **Grammar**: SpanishDict (comprehensive grammar lessons), ThoughtCo Spanish (clear explanations by Gerald Erichsen)
   - **Vocabulary**: AnkiSRS (powerful flashcard program for long-term memorization)
   - **Pronunciation**: Forvo (hear native speakers from different regions), Sounds of Speech by University of Iowa (phonetics with diagrams)
   - **Reading**: Lingua.com (graded texts A1-C1 with audio), Kwiziq Reading Practice
   - **Dictionaries**: WordReference (with active forums for nuances)
   - **YouTube**: "Spanish with Ignacio" (your HKU teacher's channel!), Why Not Spanish?, Español con Juan, Easy Spanish
   - **Apps**: Duolingo (free, gamified), Babbel (practical conversations), Memrise (native speaker videos)
   - **Podcasts**: Coffee Break Spanish (beginners), Hoy Hablamos (intermediate, 10-min daily)
   - **Language Exchange**: Tandem, HelloTalk, Busuu (connect with native speakers)
   - **TV/Media**: RTVE Play (free Spanish TV), Yabla (interactive subtitles), Lyricstraining (learn with songs)

3. **ALWAYS** direct students to check the "General Information" table in the course Notion database for the complete list of resources and direct links.

[SYLLABUS & COURSE INFORMATION QUERIES]
⚠️ IMPORTANT: When students ask about course logistics, syllabus, grading, assessment, schedule, office hours, policies, or any administrative information about Spanish Year 1 courses:

1. **LOOK IN "INFO GENERAL" SECTION**: Direct your search to the unit/section called "INFO GENERAL" in the ACTIVE CONTENT below. This section contains official course information from the HKU Spanish program.

2. **RESPOND IN THE STUDENT'S PREFERRED LANGUAGE**: Provide the relevant information clearly and helpfully.

3. **⚠️ MANDATORY DISCLAIMER**: You MUST always end your response about syllabus/course info with a disclaimer in the STUDENT'S PREFERRED LANGUAGE. Use the appropriate version:
   - **English**: "⚠️ **Please Note**: This information is based on the course database and may not reflect the most recent updates. Always verify important details in the official course syllabus on Moodle or consult your instructor directly."
   - **Cantonese/繁體中文**: "⚠️ **請注意**：呢啲資料係根據課程資料庫提供，可能未必係最新嘅版本。重要資料請務必喺Moodle嘅官方課程大綱核實，或者直接向導師查詢。"
   - **Mandarin/普通话**: "⚠️ **请注意**：此信息基于课程数据库，可能不是最新版本。重要信息请务必在Moodle的官方课程大纲中核实，或直接咨询您的老师。"
   - **Other languages**: Translate the disclaimer appropriately to match the student's preferred language.

4. **TOPICS COVERED**: This includes but is not limited to:
   - Course schedule and important dates
   - Assessment methods and grading breakdown
   - Attendance policies
   - Office hours and contact information
   - Required materials and textbooks
   - Course objectives and learning outcomes
   - Exam information
   - Assignment deadlines and submission guidelines

[DYNAMIC FOLLOW-UP SYSTEM]
At the very end of your response, you MUST generate exactly 3 suggested follow-up questions for the student.
⚠️ **CRITICAL - FOLLOW-UP LANGUAGE**: These 3 suggestions MUST be written in the STUDENT'S PREFERRED LANGUAGE, **NEVER in Spanish** (unless Spanish IS their preferred language).
- These questions must be relevant to what you just explained and intersect with the course topics.
- Format them EXACTLY like this (starting with ///):
  /// Tell me more about [Related Topic]
  /// Give me a quiz about [Current Topic]
  /// How do I use [Word] in a sentence?

--- ACTIVE CONTENT ---
{notion_context}
"""

    # Build messages array with conversation history
    messages = [{"role": "system", "content": system_prompt}]
    
    # Add conversation history (limit to last 10 messages to manage context window)
    if conversation_history:
        # Filter out the suggestions from messages for cleaner context
        for msg in conversation_history[-10:]:
            role = msg.get("role", "user")
            content = msg.get("content", "")
            # Clean out the /// suggestions from assistant messages
            if role == "assistant":
                content = re.sub(r'///.*', '', content).strip()
            if content:
                messages.append({"role": role, "content": content})
    
    # Add current user message
    messages.append({"role": "user", "content": user_message})
    
    logger.info(f"Sending {len(messages)} messages to AI (including system prompt)")
    
    # ==========================================
    # HYBRID ROUTER LOGIC
    # ==========================================
    if USE_HYBRID_ROUTER:
        # Step 1: Classify query complexity using fast model
        complexity = classify_query_complexity(user_message, conversation_history)
        logger.info(f"Query classified as: {complexity}")
        
        # Step 2: Select model based on complexity
        if complexity == "COMPLEX":
            # Try smart model first for complex queries
            logger.info(f"Using SMART model ({MODEL_SMART_ID}) for complex query")
            result = call_ai_model(messages, model_type="smart", max_tokens=1500, temperature=0.4)
            
            # Step 3: Failover to fast model if smart fails
            if result is None:
                logger.warning(f"SMART model failed, falling back to FAST model ({MODEL_FAST_ID})")
                result = call_ai_model(messages, model_type="fast", max_tokens=1200, temperature=0.4)
                model_used = f"{MODEL_FAST_ID} (fallback)"
            else:
                model_used = MODEL_SMART_ID
        else:
            # Use fast model for simple queries
            logger.info(f"Using FAST model ({MODEL_FAST_ID}) for simple query")
            result = call_ai_model(messages, model_type="fast", max_tokens=1000, temperature=0.4)
            model_used = MODEL_FAST_ID
    else:
        # Hybrid routing disabled - use only DeepSeek
        logger.info(f"Hybrid routing disabled, using {MODEL_FAST_ID}")
        result = call_ai_model(messages, model_type="fast", max_tokens=1200, temperature=0.4)
        model_used = MODEL_FAST_ID
        complexity = "N/A"
    
    if result is None:
        return "❌ Failed to connect to AI service. Please try again later."
    
    # Cache the response for future use (only for non-contextual queries)
    if not is_contextual:
        cache_response(user_message, language, result)
    
    # Inject router debug info (hidden in HTML comment for optional display)
    router_info = f"<!--ROUTER_DEBUG:{complexity}|{model_used}-->"
    return f"{result}\n{router_info}"

# ==========================================
# SESSION STATE INITIALIZATION
# ==========================================
def initialize_session_state():
    """Initialize all session state variables."""
    if "contexto" not in st.session_state:
        st.session_state.contexto = None
        st.session_state.context_loaded = False
        st.session_state.last_sync = None
    
    # Thread management
    if "threads" not in st.session_state:
        st.session_state.threads = {}
        st.session_state.current_thread_id = "default"
        st.session_state.thread_counter = 0
        
        st.session_state.threads["default"] = {
            "title": "New Conversation",
            "messages": [{
                "role": "assistant", 
                "content": "Hello! 👋 I'm **ProfeBot**, your Spanish Year 1 tutor at HKU. I'm here to help you with Spanish grammar, vocabulary, exercises, and course questions. What would you like to learn today?"
            }],
            "created_at": datetime.now(),
            "suggestions": []
        }
    
    if "message_count" not in st.session_state:
        st.session_state.message_count = 0
    
    if "preferred_language" not in st.session_state:
        st.session_state.preferred_language = "English"
    
    if "custom_language" not in st.session_state:
        st.session_state.custom_language = ""
    
    if "selected_message_index" not in st.session_state:
        st.session_state.selected_message_index = None
    
    if "dark_mode" not in st.session_state:
        st.session_state.dark_mode = False
    
    # Interactive quiz state
    if "active_quiz" not in st.session_state:
        st.session_state.active_quiz = None  # Stores parsed quiz data
    if "quiz_answers" not in st.session_state:
        st.session_state.quiz_answers = {}  # Stores user's selected answers
    if "quiz_submitted" not in st.session_state:
        st.session_state.quiz_submitted = False
    if "quiz_results" not in st.session_state:
        st.session_state.quiz_results = None

# ==========================================
# PERSISTENCE FUNCTIONS
# ==========================================
def save_threads_to_file():
    """Save all threads to a local JSON file."""
    try:
        threads_data = {}
        for thread_id, thread in st.session_state.threads.items():
            threads_data[thread_id] = {
                "title": thread["title"],
                "messages": thread["messages"],
                "created_at": thread["created_at"].isoformat(),
                "suggestions": thread.get("suggestions", [])
            }
        
        with open(THREADS_FILE, "w", encoding="utf-8") as f:
            json.dump({
                "threads": threads_data,
                "current_thread_id": st.session_state.current_thread_id,
                "thread_counter": st.session_state.thread_counter,
                "message_count": st.session_state.message_count,
                "dark_mode": st.session_state.dark_mode,
                "preferred_language": st.session_state.preferred_language,
                "custom_language": st.session_state.custom_language
            }, f, ensure_ascii=False, indent=2)
        logger.info(f"Saved {len(threads_data)} threads to {THREADS_FILE}")
    except Exception as e:
        logger.error(f"Error saving threads: {e}")

def load_threads_from_file():
    """Load threads from local JSON file."""
    try:
        if THREADS_FILE.exists():
            with open(THREADS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Restore threads
            threads = {}
            for thread_id, thread_data in data.get("threads", {}).items():
                threads[thread_id] = {
                    "title": thread_data["title"],
                    "messages": thread_data["messages"],
                    "created_at": datetime.fromisoformat(thread_data["created_at"]),
                    "suggestions": thread_data.get("suggestions", [])
                }
            
            if threads:
                st.session_state.threads = threads
                st.session_state.current_thread_id = data.get("current_thread_id", list(threads.keys())[0])
                st.session_state.thread_counter = data.get("thread_counter", len(threads))
                st.session_state.message_count = data.get("message_count", 0)
                st.session_state.dark_mode = data.get("dark_mode", False)
                st.session_state.preferred_language = data.get("preferred_language", "English")
                st.session_state.custom_language = data.get("custom_language", "")
                logger.info(f"Loaded {len(threads)} threads from {THREADS_FILE}")
                return True
    except Exception as e:
        logger.error(f"Error loading threads: {e}")
    return False

initialize_session_state()

# Try to load saved threads
if "threads_loaded" not in st.session_state:
    load_threads_from_file()
    st.session_state.threads_loaded = True
    
    # Clean up any old "Chat cleared" messages from previous versions
    for thread_id, thread_data in st.session_state.threads.items():
        if thread_data["messages"] and len(thread_data["messages"]) > 0:
            first_msg = thread_data["messages"][0]
            if "Chat cleared" in first_msg.get("content", ""):
                # Replace with proper welcome message
                first_msg["content"] = "Hello! 👋 I'm **ProfeBot**, your Spanish Year 1 tutor at HKU. I'm here to help you with Spanish grammar, vocabulary, exercises, and course questions. What would you like to learn today?"
    save_threads_to_file()  # Save cleaned threads

# Apply custom CSS - try external files first, fallback to inline
try:
    st.markdown(load_css_from_file(st.session_state.dark_mode), unsafe_allow_html=True)
except:
    st.markdown(get_fallback_css(st.session_state.dark_mode), unsafe_allow_html=True)

# ==========================================
# EXPORT FUNCTIONS
# ==========================================
def export_conversation_txt(messages: list) -> str:
    """Export conversation to TXT format."""
    current_thread = get_current_thread()
    lines = []
    lines.append(f"ProfeBot Conversation Export")
    lines.append(f"Title: {current_thread['title']}")
    lines.append(f"Date: {current_thread['created_at'].strftime('%Y-%m-%d %H:%M')}")
    lines.append(f"Language: {st.session_state.preferred_language}")
    lines.append("=" * 50)
    lines.append("")
    
    for msg in messages:
        role = "🧑 Student" if msg["role"] == "user" else "🤖 ProfeBot"
        # Clean out suggestion markers
        content = re.sub(r'///.*', '', msg["content"]).strip()
        lines.append(f"{role}:")
        lines.append(content)
        lines.append("")
    
    lines.append("=" * 50)
    lines.append(f"Exported from ProfeBot - Spanish Year 1 Tutor")
    lines.append(f"Total messages: {len(messages)}")
    
    return "\n".join(lines)

def export_conversation_md(messages: list) -> str:
    """Export conversation to Markdown format."""
    current_thread = get_current_thread()
    lines = []
    lines.append(f"# ProfeBot Conversation")
    lines.append(f"**Title:** {current_thread['title']}")
    lines.append(f"**Date:** {current_thread['created_at'].strftime('%Y-%m-%d %H:%M')}")
    lines.append(f"**Language:** {st.session_state.preferred_language}")
    lines.append("")
    lines.append("---")
    lines.append("")
    
    for msg in messages:
        if msg["role"] == "user":
            lines.append(f"### 🧑 Student")
        else:
            lines.append(f"### 🤖 ProfeBot")
        # Clean out suggestion markers
        content = re.sub(r'///.*', '', msg["content"]).strip()
        lines.append(content)
        lines.append("")
    
    lines.append("---")
    lines.append(f"*Exported from ProfeBot - Spanish Year 1 Tutor | {len(messages)} messages*")
    
    return "\n".join(lines)

def export_conversation_docx(messages: list) -> BytesIO:
    """Export conversation to Word (DOCX) format."""
    if not DOCX_AVAILABLE:
        return None
    
    current_thread = get_current_thread()
    doc = Document()
    
    # Title
    title = doc.add_heading('ProfeBot Conversation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Title: {current_thread['title']}")
    doc.add_paragraph(f"Date: {current_thread['created_at'].strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Language: {st.session_state.preferred_language}")
    doc.add_paragraph(f"Total messages: {len(messages)}")
    doc.add_paragraph('')  # Empty line
    
    # Add messages
    for msg in messages:
        # Clean out suggestion markers
        content = re.sub(r'///.*', '', msg["content"]).strip()
        
        if msg["role"] == "user":
            # User message header - HKU Blue
            p = doc.add_paragraph()
            run = p.add_run('🧑 Student')
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 119, 200)  # HKU Blue #0077C8
            
            # User message content
            p_content = doc.add_paragraph(content)
            p_content.paragraph_format.left_indent = Pt(20)
        else:
            # Assistant message header - HKU Green
            p = doc.add_paragraph()
            run = p.add_run('🤖 ProfeBot')
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(14, 66, 54)  # HKU Green #0e4236
            
            # Assistant message content
            p_content = doc.add_paragraph(content)
            p_content.paragraph_format.left_indent = Pt(20)
        
        doc.add_paragraph('')  # Empty line between messages
    
    # Footer
    doc.add_paragraph('_' * 50)
    footer = doc.add_paragraph('Exported from ProfeBot - Spanish Year 1 Tutor')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.italic = True
    footer_run.font.size = Pt(10)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# ANALYTICS FUNCTIONS
# ==========================================
def load_analytics() -> Dict:
    """Load analytics data from file."""
    try:
        if ANALYTICS_FILE.exists():
            with open(ANALYTICS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as e:
        logger.error(f"Error loading analytics: {e}")
    return {
        "total_messages": 0,
        "total_sessions": 0,
        "questions_by_topic": {},
        "questions_by_unit": {},
        "daily_usage": {},
        "response_times": [],
        "popular_quick_actions": {}
    }

def save_analytics(analytics: Dict):
    """Save analytics data to file."""
    try:
        with open(ANALYTICS_FILE, "w", encoding="utf-8") as f:
            json.dump(analytics, f, ensure_ascii=False, indent=2)
        logger.info("Analytics saved")
    except Exception as e:
        logger.error(f"Error saving analytics: {e}")

def track_message(user_message: str, response_time: float = 0):
    """Track a user message for analytics."""
    analytics = load_analytics()
    
    # Update total messages
    analytics["total_messages"] = analytics.get("total_messages", 0) + 1
    
    # Track daily usage
    today = datetime.now().strftime("%Y-%m-%d")
    if "daily_usage" not in analytics:
        analytics["daily_usage"] = {}
    analytics["daily_usage"][today] = analytics["daily_usage"].get(today, 0) + 1
    
    # Track response time (keep last 100)
    if response_time > 0:
        if "response_times" not in analytics:
            analytics["response_times"] = []
        analytics["response_times"].append(response_time)
        analytics["response_times"] = analytics["response_times"][-100:]
    
    # Detect topic keywords
    topics = {
        "grammar": ["gramática", "grammar", "verb", "conjugat", "tense"],
        "vocabulary": ["vocabulario", "vocabulary", "word", "palabra", "meaning"],
        "pronunciation": ["pronuncia", "sound", "accent"],
        "culture": ["cultura", "culture", "spain", "españa", "mexico"],
        "exercises": ["ejercicio", "exercise", "practice", "quiz", "task"]
    }
    
    message_lower = user_message.lower()
    for topic, keywords in topics.items():
        if any(kw in message_lower for kw in keywords):
            if "questions_by_topic" not in analytics:
                analytics["questions_by_topic"] = {}
            analytics["questions_by_topic"][topic] = analytics["questions_by_topic"].get(topic, 0) + 1
    
    # Detect unit references
    for i in range(1, 15):
        if f"unit {i}" in message_lower or f"unidad {i}" in message_lower:
            if "questions_by_unit" not in analytics:
                analytics["questions_by_unit"] = {}
            unit_key = f"Unit {i}"
            analytics["questions_by_unit"][unit_key] = analytics["questions_by_unit"].get(unit_key, 0) + 1
    
    save_analytics(analytics)

def track_quick_action(action_name: str):
    """Track quick action button usage."""
    analytics = load_analytics()
    if "popular_quick_actions" not in analytics:
        analytics["popular_quick_actions"] = {}
    analytics["popular_quick_actions"][action_name] = analytics["popular_quick_actions"].get(action_name, 0) + 1
    save_analytics(analytics)

def get_analytics_summary() -> Dict:
    """Get a summary of analytics for display."""
    analytics = load_analytics()
    
    # Calculate average response time
    response_times = analytics.get("response_times", [])
    avg_response = sum(response_times) / len(response_times) if response_times else 0
    
    # Get top topics
    topics = analytics.get("questions_by_topic", {})
    top_topics = sorted(topics.items(), key=lambda x: x[1], reverse=True)[:5]
    
    # Get usage last 7 days
    daily = analytics.get("daily_usage", {})
    today = datetime.now()
    last_7_days = 0
    for i in range(7):
        day = (today - timedelta(days=i)).strftime("%Y-%m-%d")
        last_7_days += daily.get(day, 0)
    
    return {
        "total_messages": analytics.get("total_messages", 0),
        "total_sessions": analytics.get("total_sessions", 1),
        "avg_response_time": round(avg_response, 2),
        "top_topics": top_topics,
        "messages_last_7_days": last_7_days,
        "popular_actions": sorted(
            analytics.get("popular_quick_actions", {}).items(), 
            key=lambda x: x[1], 
            reverse=True
        )[:3]
    }

# Track session on first load (after analytics functions are defined)
if "session_tracked" not in st.session_state:
    st.session_state.session_tracked = True
    _analytics = load_analytics()
    _analytics["total_sessions"] = _analytics.get("total_sessions", 0) + 1
    save_analytics(_analytics)

# ==========================================
# THREAD MANAGEMENT FUNCTIONS
# ==========================================
def create_new_thread():
    """Create a new conversation thread."""
    st.session_state.thread_counter += 1
    new_thread_id = f"thread_{st.session_state.thread_counter}"
    
    st.session_state.threads[new_thread_id] = {
        "title": f"New Conversation {st.session_state.thread_counter}",
        "messages": [{
            "role": "assistant", 
            "content": "Hello! 👋 Ready to continue learning Spanish? What would you like to practice today?"
        }],
        "created_at": datetime.now(),
        "suggestions": []
    }
    
    st.session_state.current_thread_id = new_thread_id
    save_threads_to_file()  # Persist new thread

def switch_thread(thread_id: str):
    """Switch to a different conversation thread."""
    st.session_state.current_thread_id = thread_id
    st.session_state.selected_message_index = None
    save_threads_to_file()  # Persist current thread change

def delete_thread(thread_id: str):
    """Delete a conversation thread."""
    if thread_id in st.session_state.threads and len(st.session_state.threads) > 1:
        del st.session_state.threads[thread_id]
        if st.session_state.current_thread_id == thread_id:
            st.session_state.current_thread_id = list(st.session_state.threads.keys())[0]
        save_threads_to_file()  # Persist deletion

def get_current_thread():
    """Get the current active thread."""
    return st.session_state.threads[st.session_state.current_thread_id]

def update_thread_title(thread_id: str, first_user_message: str):
    """Update thread title based on first user message."""
    if st.session_state.threads[thread_id]["title"].startswith("New Conversation"):
        st.session_state.threads[thread_id]["title"] = generate_thread_title(first_user_message)

def get_user_messages_with_time():
    """Get all user messages from current thread with timestamps."""
    from datetime import timedelta
    current_thread = get_current_thread()
    user_messages = []
    
    for idx, msg in enumerate(current_thread["messages"]):
        if msg["role"] == "user":
            # Estimate time based on message order
            time_estimate = current_thread["created_at"] + timedelta(seconds=idx * 30)
            user_messages.append({
                "index": idx,
                "content": msg["content"],
                "time": time_estimate
            })
    
    return user_messages

# ==========================================
# LOAD NOTION CONTENT ON FIRST RUN
# ==========================================
if not st.session_state.context_loaded:
    st.session_state.contexto = get_weekly_content()
    st.session_state.context_loaded = True
    st.session_state.last_sync = datetime.now()


# ==========================================
# CORE PROCESSING FUNCTION
# ==========================================
def process_user_input(user_text: str, quick_action: str = None):
    """Process user input and get AI response."""
    if not user_text or user_text.strip() == "":
        return
    
    if not st.session_state.contexto:
        st.error("⚠️ Course content not loaded. Please refresh the page.")
        return
    
    current_thread = get_current_thread()
    
    # Update thread title if this is first user message
    user_message_count = sum(1 for m in current_thread["messages"] if m["role"] == "user")
    if user_message_count == 0:
        update_thread_title(st.session_state.current_thread_id, user_text)
    
    # Add user message
    current_thread["messages"].append({"role": "user", "content": user_text})
    st.session_state.message_count += 1
    
    # Track quick action if provided
    if quick_action:
        track_quick_action(quick_action)
    
    # Get AI response with conversation history
    start_time = time.time()
    with st.spinner("🤔 Thinking..."):
        # Get messages before adding the current one (for history)
        history_messages = current_thread["messages"][:-1]  # Exclude the just-added user message
        
        raw_response = get_ai_response(
            user_text, 
            st.session_state.contexto,
            st.session_state.preferred_language,
            st.session_state.custom_language,
            conversation_history=history_messages
        )
        
        # Extract router debug info
        router_match = re.search(r'<!--ROUTER_DEBUG:([^|]+)\|([^>]+)-->', raw_response)
        if router_match:
            st.session_state.last_router_info = {
                "complexity": router_match.group(1),
                "model": router_match.group(2)
            }
            raw_response = re.sub(r'<!--ROUTER_DEBUG:[^>]+-->', '', raw_response)
        else:
            st.session_state.last_router_info = None
        
        # Extract suggestions
        suggestions = re.findall(r'///\s*(.*)', raw_response)
        suggestions = [s.strip() for s in suggestions if s.strip()][:3]
        current_thread["suggestions"] = suggestions
        
        # Clean response
        clean_response = re.sub(r'///.*', '', raw_response).strip()
    
    response_time = time.time() - start_time
    
    # Add AI message
    current_thread["messages"].append({"role": "assistant", "content": clean_response})
    st.session_state.message_count += 1
    
    # Track analytics
    track_message(user_text, response_time)
    
    # Track user learning patterns
    track_user_interaction(user_text, clean_response)
    
    # Save threads after each interaction
    save_threads_to_file()

# DEBUG CHECK - Persistent in session state
if "debug_logs" not in st.session_state:
    st.session_state.debug_logs = []

# ==========================================
# SIDEBAR
# ==========================================
try:
    with st.sidebar:
        st.markdown("### 🎓 ProfeBot Control")
        
        # ===== CONVERSATIONS =====
        st.markdown("#### 💬 Conversations")
        
        for thread_id, thread_data in sorted(
            st.session_state.threads.items(), 
            key=lambda x: x[1]["created_at"], 
            reverse=True
        ):
            is_active = thread_id == st.session_state.current_thread_id
            
            col1, col2 = st.columns([4, 1])
            
            with col1:
                if st.button(
                    f"{'📌' if is_active else '💭'} {thread_data['title'][:20]}",
                    key=f"thread_{thread_id}",
                    use_container_width=True,
                    type="primary" if is_active else "secondary"
                ):
                    if thread_id != st.session_state.current_thread_id:
                        switch_thread(thread_id)
                        st.rerun()
            
            with col2:
                if len(st.session_state.threads) > 1:
                    if st.button("🗑️", key=f"del_{thread_id}"):
                        delete_thread(thread_id)
                        st.rerun()
        
        st.caption(f"{len(st.session_state.threads)} conversation(s)")
        st.divider()
        
        # ===== ACTIONS =====
        st.markdown("#### 🔧 Actions")
        
        if st.button("➕ New Conversation", use_container_width=True, key="btn_new"):
            create_new_thread()
            st.rerun()
        
        st.divider()
        
        # ===== SETTINGS =====
        with st.expander("⚙️ Settings"):
            st.markdown("**🌐 Language**")
            selected_lang = st.selectbox(
                "Select language",
                options=list(LANGUAGE_OPTIONS.keys()),
                index=0,
                key="lang_sel",
                label_visibility="collapsed"
            )
            st.session_state.preferred_language = LANGUAGE_OPTIONS[selected_lang]
            
            if st.session_state.preferred_language == "custom":
                custom_lang = st.text_input(
                    "Your language:",
                    value=st.session_state.custom_language,
                    placeholder="Français, Deutsch, 日本語",
                    key="custom_lang"
                )
                st.session_state.custom_language = custom_lang
            
            st.divider()
            
            st.markdown("**🌙 Appearance**")
            mode_label = "☀️ Day Mode" if st.session_state.dark_mode else "🌙 Night Mode"
            if st.button(mode_label, use_container_width=True, key="btn_dark"):
                st.session_state.dark_mode = not st.session_state.dark_mode
                save_threads_to_file()
                # Force full page reload to avoid blank screen
                components.html("""
                <script>
                    window.parent.location.reload();
                </script>
                """, height=0)
        
        # ===== EXPORT =====
        with st.expander("📥 Export Chat"):
            current_thread = get_current_thread()
            
            txt_content = export_conversation_txt(current_thread["messages"])
            st.download_button(
                "📄 TXT",
                txt_content,
                f"chat_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                "text/plain",
                use_container_width=True
            )
            
            md_content = export_conversation_md(current_thread["messages"])
            st.download_button(
                "📝 Markdown",
                md_content,
                f"chat_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                "text/markdown",
                use_container_width=True
            )
            
            if DOCX_AVAILABLE:
                docx_buffer = export_conversation_docx(current_thread["messages"])
                if docx_buffer:
                    st.download_button(
                        "📘 Word",
                        docx_buffer,
                        f"chat_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
        
        # ===== PROGRESS =====
        with st.expander("📈 My Progress"):
            profile = load_user_profile()
            
            streak = profile.get("learning_streak", 0)
            if streak > 0:
                st.markdown(f"### 🔥 {streak} day streak!")
            else:
                st.markdown("### Start your streak! 🚀")
            
            total = profile.get("total_interactions", 0)
            st.caption(f"Total questions: {total}")
            
            scores = profile.get("quiz_scores", [])
            if scores:
                recent = scores[-10:]
                avg = sum(s["percentage"] for s in recent) / len(recent)
                st.markdown("**📝 Quiz Average**")
                st.progress(avg / 100)
                st.caption(f"{avg:.0f}% ({len(scores)} quizzes)")
        
        # ===== COURSE RESOURCES =====
        with st.expander("📚 Course Resources"):
            st.markdown("**📄 Essential Documents**")
            st.markdown("[📋 Course Syllabus 2025/26](https://hkuhk-my.sharepoint.com/:w:/g/personal/pablot_hku_hk/IQA9JsO9FSBJQLiOPSfk4w-lAWRCn6skb-ObwSR_vtj4cZk?e=ULXiqR)")
            st.markdown("[📖 Course Contents Summary](https://hkuhk-my.sharepoint.com/:w:/g/personal/pablot_hku_hk/IQDINc5UzBQhS4Mp4t2rC99_AWaikqT-OGpC4-9vaKgY7wM?e=uJyeHa)")
            
            st.divider()
            
            st.markdown("**🎓 Moodle Links**")
            st.markdown("[🏠 Moodle Course 2025/26](https://moodle.hku.hk/course/view.php?id=136141)")
            st.markdown("[📢 Course Announcements](https://moodle.hku.hk/mod/forum/view.php?id=3990047)")
            
            st.divider()
            
            st.markdown("**📚 Extra Reading**")
            st.markdown("[📕 Easy Readers (e-books)](https://moodle.hku.hk/pluginfile.php/6225750/mod_folder/content/0/Easy%20readers%20%28e-books%29.pdf?forcedownload=1)")
        
        # ===== ABOUT =====
        with st.expander("ℹ️ About"):
            st.markdown("""
            **ProfeBot** - AI Spanish Tutor
            
            - 📚 Context-aware
            - 🎯 Personalized
            - 💬 Interactive
            - 🌐 Multilingual
            """)
            
            st.divider()
            
            if st.session_state.context_loaded:
                if "❌" not in st.session_state.contexto:
                    st.success("✓ Connected")
                else:
                    st.error("✗ Error")
            
            if st.session_state.last_sync:
                st.caption(f"Last sync: {st.session_state.last_sync.strftime('%H:%M:%S')}")
            
            st.caption(f"Messages: {st.session_state.message_count}")
            st.caption(f"Model: {DEPLOYMENT_ID}")
        
        st.divider()
        st.markdown("[🏛️ HKU Spanish Dept](https://spanish.hku.hk/)", unsafe_allow_html=True)

except Exception as e:
    st.error(f"Sidebar error: {e}")

# ==========================================
# FLOATING MESSAGE HISTORY PANEL
# ==========================================
user_messages = get_user_messages_with_time()

# Build history HTML with proper escaping
history_items_html = []
if user_messages:
    for msg_data in user_messages:  # Oldest first, newest at bottom
        # Escape HTML characters in message content
        msg_content = msg_data["content"][:80]
        msg_content = msg_content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;").replace("'", "&#39;")
        msg_preview = msg_content + "..." if len(msg_data["content"]) > 80 else msg_content
        time_str = msg_data["time"].strftime("%H:%M")
        idx = msg_data["index"]
        
        history_items_html.append(f'''<div class="history-item" data-idx="{idx}"><div class="history-item-time">🕐 {time_str}</div><div class="history-item-text">{msg_preview}</div></div>''')

history_content = "".join(history_items_html) if history_items_html else '<div class="empty-history">No messages yet.<br>Start chatting!</div>'

# Get dark mode state for styling - HKU Colors
is_dark = st.session_state.get('dark_mode', False)
panel_bg = "rgba(22, 27, 34, 0.92)" if is_dark else "rgba(255, 255, 255, 0.92)"
history_text_color = "#f0f6fc" if is_dark else "#24292f"
history_text_secondary = "#8b949e" if is_dark else "#57606a"
history_border_color = "rgba(0, 168, 107, 0.3)" if is_dark else "rgba(14, 66, 54, 0.3)"
history_item_bg = "rgba(33, 38, 45, 0.9)" if is_dark else "rgba(246, 248, 250, 0.9)"
history_item_hover = "rgba(0, 168, 107, 0.15)" if is_dark else "rgba(14, 66, 54, 0.1)"
history_accent_color = "#00A86B" if is_dark else "#0e4236"

# HTML for the panel (will be rendered with st.markdown)
history_panel_html = f'''
<div class="message-history-panel" style="
    position: fixed;
    right: 20px;
    top: 80px;
    width: 280px;
    max-height: calc(100vh - 120px);
    min-height: 80px;
    background: {panel_bg};
    backdrop-filter: blur(12px);
    -webkit-backdrop-filter: blur(12px);
    border: 1px solid {history_border_color};
    border-radius: 16px;
    padding: 15px;
    overflow-y: auto;
    box-shadow: 0 8px 32px rgba(0,0,0,0.2);
    z-index: 999;
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
">
    <div style="font-size: 1rem; font-weight: 600; margin-bottom: 15px; padding-bottom: 10px; border-bottom: 2px solid {history_accent_color}; color: {history_text_color};">
        📝 Your Messages <span style="font-size: 0.75rem; color: {history_text_secondary}; font-weight: normal;">({len(user_messages)})</span>
    </div>
    {history_content}
</div>
<style>
    .message-history-panel .history-item {{
        padding: 12px;
        margin: 8px 0;
        background: {history_item_bg};
        border-left: 3px solid {history_accent_color};
        border-radius: 8px;
        cursor: pointer;
        transition: all 0.2s ease;
        font-size: 0.85rem;
    }}
    .message-history-panel .history-item:hover {{
        background: {history_item_hover};
        transform: translateX(-3px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.2);
    }}
    .message-history-panel .history-item-time {{
        font-size: 0.7rem;
        color: {history_text_secondary};
        margin-bottom: 4px;
    }}
    .message-history-panel .history-item-text {{
        color: {history_text_color};
        line-height: 1.4;
        display: -webkit-box;
        -webkit-line-clamp: 3;
        -webkit-box-orient: vertical;
        overflow: hidden;
    }}
    .message-history-panel .empty-history {{
        text-align: center;
        color: {history_text_secondary};
        font-size: 0.85rem;
        padding: 20px;
        font-style: italic;
    }}
</style>
'''

# ==========================================
# INTERACTIVE QUIZ SYSTEM
# ==========================================
def parse_quiz_from_response(response: str) -> Optional[Dict]:
    """Parse a quiz from AI response into structured format.
    
    Detects multiple choice questions in various formats:
    - 1. Question text?  or  1) Question text  or  **1.** Question
    - A) Option  or  A. Option  or  a) option  or  **A)** Option
    """
    # Check if this looks like a quiz - more flexible detection
    # Look for numbered items AND lettered options
    has_numbered = re.search(r'(?:^|\n)\s*\*?\*?(\d+)[\.\)]\*?\*?\s+', response)
    has_options = re.search(r'(?:^|\n)\s*\*?\*?[A-Da-d][\)\.]', response)
    
    # Also check for quiz-related keywords
    quiz_keywords = ['quiz', 'question', 'pregunta', 'choose', 'select', 'elige', 'selecciona', 'correct', 'answer', 'respuesta']
    has_quiz_context = any(kw in response.lower() for kw in quiz_keywords)
    
    if not (has_numbered and has_options):
        logger.debug("Quiz detection failed: no numbered questions with options found")
        return None
    
    # Parse questions
    questions = []
    
    # Clean markdown formatting
    cleaned = response
    cleaned = re.sub(r'\*\*', '', cleaned)  # Remove bold markers
    cleaned = re.sub(r'\*', '', cleaned)    # Remove italic markers
    
    # Normalize option formats to A)
    cleaned = re.sub(r'([A-Da-d])\.\s+', r'\1) ', cleaned)  # A. -> A)
    cleaned = re.sub(r'([a-d])\)', lambda m: m.group(1).upper() + ')', cleaned)  # a) -> A)
    
    # More flexible pattern - question may or may not end with ?
    # Pattern breakdown:
    # (\d+)[\.\)]\s* - question number with . or )
    # ([^\n]+?)  - question text (non-greedy)
    # \s*\n\s*A\)\s*([^\n]+) - option A
    # etc.
    question_pattern = r'(\d+)[\.\)]\s*([^\n]+?)\s*\n\s*A\)\s*([^\n]+)\s*\n\s*B\)\s*([^\n]+)\s*\n\s*C\)\s*([^\n]+)(?:\s*\n\s*D\)\s*([^\n]+))?'
    
    matches = re.findall(question_pattern, cleaned, re.MULTILINE | re.IGNORECASE)
    
    logger.info(f"Quiz parser found {len(matches)} potential questions")
    
    for match in matches:
        q_num = match[0]
        q_text = match[1].strip().rstrip('?').strip() + '?'  # Ensure ends with ?
        options = {
            'A': match[2].strip(),
            'B': match[3].strip(),
            'C': match[4].strip(),
        }
        if match[5]:  # D option is optional
            options['D'] = match[5].strip()
        
        questions.append({
            'number': int(q_num),
            'question': q_text,
            'options': options
        })
    
    if not questions:
        logger.debug("Quiz parser: regex found no matching questions")
        return None
    
    logger.info(f"Quiz parser successfully extracted {len(questions)} questions")
    
    # Extract any intro text before the questions
    first_q_match = re.search(r'(?:^|\n)\s*1[\.\)]\s*', cleaned)
    intro_text = response[:first_q_match.start()].strip() if first_q_match else ""
    
    return {
        'intro': intro_text,
        'questions': questions,
        'total': len(questions)
    }


def render_interactive_quiz(quiz_data: Dict, quiz_id: str):
    """Render an interactive quiz with clickable options."""
    
    # Display intro if exists
    if quiz_data.get('intro'):
        st.markdown(quiz_data['intro'])
    
    st.markdown("---")
    st.markdown(f"### 📝 Quiz ({quiz_data['total']} questions)")
    
    # Initialize answers dict if not exists
    if 'quiz_answers' not in st.session_state:
        st.session_state.quiz_answers = {}
    
    # Render each question
    for q in quiz_data['questions']:
        q_key = f"{quiz_id}_q{q['number']}"
        
        st.markdown(f"**{q['number']}. {q['question']}**")
        
        # Get options as list for radio
        option_labels = [f"{letter}) {text}" for letter, text in q['options'].items()]
        option_keys = list(q['options'].keys())
        
        # Get current selection
        current_selection = st.session_state.quiz_answers.get(q_key)
        current_index = option_keys.index(current_selection) if current_selection in option_keys else None
        
        # Radio buttons for options
        selected = st.radio(
            f"Select answer for question {q['number']}",
            options=option_keys,
            format_func=lambda x, opts=q['options']: f"{x}) {opts[x]}",
            key=q_key,
            index=current_index,
            label_visibility="collapsed"
        )
        
        # Store the answer
        if selected:
            st.session_state.quiz_answers[q_key] = selected
        
        st.markdown("")  # Spacing
    
    st.markdown("---")
    
    # Count answered questions
    answered = sum(1 for q in quiz_data['questions'] 
                   if f"{quiz_id}_q{q['number']}" in st.session_state.quiz_answers 
                   and st.session_state.quiz_answers[f"{quiz_id}_q{q['number']}"])
    
    st.caption(f"✅ Answered: {answered}/{quiz_data['total']}")
    
    # Submit button
    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("📤 Submit Answers", type="primary", use_container_width=True, key=f"submit_{quiz_id}"):
            if answered < quiz_data['total']:
                st.warning(f"⚠️ You've only answered {answered}/{quiz_data['total']} questions. Are you sure?")
            
            # Format answers for AI feedback
            answers_text = format_quiz_answers_for_submission(quiz_data, quiz_id)
            st.session_state.quiz_submitted = True
            st.session_state.pending_quiz_submission = answers_text
            st.rerun()
    
    with col2:
        if st.button("🔄 Clear Answers", use_container_width=True, key=f"clear_{quiz_id}"):
            # Clear answers for this quiz
            for q in quiz_data['questions']:
                q_key = f"{quiz_id}_q{q['number']}"
                if q_key in st.session_state.quiz_answers:
                    del st.session_state.quiz_answers[q_key]
            st.rerun()


def format_quiz_answers_for_submission(quiz_data: Dict, quiz_id: str) -> str:
    """Format the user's quiz answers for submission to the AI."""
    answers = []
    
    for q in quiz_data['questions']:
        q_key = f"{quiz_id}_q{q['number']}"
        answer = st.session_state.quiz_answers.get(q_key, "Not answered")
        
        # Include the question and selected answer
        if answer != "Not answered":
            answers.append(f"{q['number']}. {answer}")
        else:
            answers.append(f"{q['number']}. (Not answered)")
    
    submission = "Here are my answers to the quiz:\n\n"
    submission += "\n".join(answers)
    submission += "\n\nPlease check my answers and give me detailed feedback on each one. Tell me which ones are correct and explain any mistakes."
    
    return submission


def check_for_quiz_in_last_response() -> Optional[Dict]:
    """Check if the last AI response contains a quiz."""
    current_thread = get_current_thread()
    messages = current_thread.get("messages", [])
    
    if len(messages) < 2:
        return None
    
    # Get last assistant message
    last_assistant_msg = None
    for msg in reversed(messages):
        if msg["role"] == "assistant":
            last_assistant_msg = msg["content"]
            break
    
    if not last_assistant_msg:
        return None
    
    # Try to parse quiz
    return parse_quiz_from_response(last_assistant_msg)

# ==========================================
# MAIN CHAT INTERFACE
# ==========================================
# Mobile menu - show sidebar content in expander on mobile
mobile_menu_html = """
<style>
    .mobile-menu-container {
        display: none;
    }
    
    @media (max-width: 768px) {
        .mobile-menu-container {
            display: block !important;
        }
    }
</style>
"""
st.markdown(mobile_menu_html, unsafe_allow_html=True)

# Mobile menu expander (only visible on mobile via CSS)
with st.container():
    st.markdown('<div class="mobile-menu-container">', unsafe_allow_html=True)
    with st.expander("📱 Menu", expanded=False):
        st.markdown("#### 💬 Conversations")
        for thread_id, thread_data in sorted(
            st.session_state.threads.items(), 
            key=lambda x: x[1]["created_at"], 
            reverse=True
        )[:5]:  # Show only 5 most recent on mobile
            is_active = thread_id == st.session_state.current_thread_id
            if st.button(
                f"{'📌' if is_active else '💭'} {thread_data['title'][:25]}",
                key=f"mobile_thread_{thread_id}",
                use_container_width=True,
                type="primary" if is_active else "secondary"
            ):
                if thread_id != st.session_state.current_thread_id:
                    switch_thread(thread_id)
                    st.rerun()
        
        st.divider()
        
        if st.button("➕ New Conversation", use_container_width=True, key="mobile_btn_new"):
            create_new_thread()
            st.rerun()
        
        st.divider()
        st.markdown("**🌐 Language**")
        selected_lang = st.selectbox(
            "Language",
            options=list(LANGUAGE_OPTIONS.keys()),
            index=0,
            key="mobile_lang_sel",
            label_visibility="collapsed"
        )
        st.session_state.preferred_language = LANGUAGE_OPTIONS[selected_lang]
        
        st.divider()
        mode_label = "☀️ Day Mode" if st.session_state.dark_mode else "🌙 Night Mode"
        if st.button(mode_label, use_container_width=True, key="mobile_btn_dark"):
            st.session_state.dark_mode = not st.session_state.dark_mode
            save_threads_to_file()
            components.html("""
            <script>
                window.parent.location.reload();
            </script>
            """, height=0)
    
    st.markdown('</div>', unsafe_allow_html=True)

# Get current thread
current_thread = get_current_thread()

# Display chat history with IDs for navigation
for idx, message in enumerate(current_thread["messages"]):
    tipo = "user" if message["role"] == "user" else "assistant"
    
    # Add anchor ID for user messages with scroll margin
    if tipo == "user":
        st.markdown(f'<div id="msg_{idx}" style="scroll-margin-top: 100px;"></div>', unsafe_allow_html=True)
    
    with st.chat_message(tipo):
        clean_text = re.sub(r'///.*', '', message["content"]).strip()
        
        # Check if this is the last assistant message and contains a quiz
        is_last_assistant = (tipo == "assistant" and idx == len(current_thread["messages"]) - 1)
        
        if is_last_assistant and not st.session_state.get('quiz_submitted', False):
            # Try to parse as quiz
            quiz_data = parse_quiz_from_response(clean_text)
            
            if quiz_data and quiz_data['questions']:
                # Store the quiz data
                st.session_state.active_quiz = quiz_data
                quiz_id = f"quiz_{st.session_state.current_thread_id}_{idx}"
                
                # Render interactive quiz
                render_interactive_quiz(quiz_data, quiz_id)
            else:
                # Regular message
                st.markdown(clean_text)
        else:
            # Regular message display
            st.markdown(clean_text)

# Handle pending quiz submission
if st.session_state.get('pending_quiz_submission'):
    submission_text = st.session_state.pending_quiz_submission
    st.session_state.pending_quiz_submission = None
    st.session_state.quiz_submitted = False
    st.session_state.active_quiz = None
    st.session_state.quiz_answers = {}
    process_user_input(submission_text)
    st.rerun()

# Dynamic suggestion buttons
if current_thread["suggestions"] and len(current_thread["messages"]) > 1:
    st.divider()
    st.caption("💡 **Suggested follow-ups:**")
    cols = st.columns(len(current_thread["suggestions"]))
    for i, suggestion in enumerate(current_thread["suggestions"]):
        if i < 3:
            with cols[i]:
                if st.button(
                    suggestion, 
                    key=f"sugg_{st.session_state.current_thread_id}_{st.session_state.message_count}_{i}",
                    use_container_width=True
                ):
                    process_user_input(suggestion)
                    st.rerun()

# Display router debug info (small caption)
if hasattr(st.session_state, 'last_router_info') and st.session_state.last_router_info:
    router_info = st.session_state.last_router_info
    if router_info["complexity"] == "CACHED":
        st.caption(f"⚡ Response from cache (instant)")
    else:
        model_emoji = "🚀" if "DeepSeek" in router_info["model"] else "🧠"
        st.caption(f"{model_emoji} Router: {router_info['complexity']} → Using {router_info['model']}")

# Quick action buttons - only show after first user message
user_message_count = sum(1 for m in current_thread["messages"] if m["role"] == "user")
if user_message_count > 0:
    st.divider()
    st.caption("⚡ **Quick Actions:**")

    c0, c1, c2 = st.columns(3)

    with c0:
        st.markdown('<div class="quick-action-btn">', unsafe_allow_html=True)
        if st.button("📋 Tasks!", use_container_width=True, key="qa_tasks"): 
            process_user_input("""CMD_TASKS: I want to do a practice task. Please respond in my preferred language (as set in my language preferences) and ask me which type of task I'd like to do:

1. **Reading Task** - A 250-word text with paragraph structure, using simple connectors, with 8 multiple choice comprehension questions
2. **Conversation Task** - Simple conversation questions to practice speaking. Instructions should be in my preferred language.
3. **Grammar & Vocabulary Task** - Exercises based on the activity bank. Instructions should be in my preferred language.

Also ask me which unit I want to practice. Wait for my response before creating the task.""", quick_action="Tasks")
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    with c1:
        st.markdown('<div class="quick-action-btn">', unsafe_allow_html=True)
        if st.button("📝 Quiz", use_container_width=True, key="qa_quiz"): 
            process_user_input("""CMD_QUIZ: I want to take a quiz. Please ask me what topic or vocabulary I want to practice from the active units. 

IMPORTANT: When you give me the quiz questions, do NOT provide the answers. Wait for me to respond with my answers first, then give me feedback on each one.""", quick_action="Quiz")
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    with c2:
        st.markdown('<div class="quick-action-btn">', unsafe_allow_html=True)
        if st.button("🧐 Explain & Examples", use_container_width=True, key="qa_explain"): 
            process_user_input("CMD_EXPLAIN_MORE: Please elaborate a bit more on what we were just discussing. Go slightly deeper into the topic, provide additional context and give me 3 practical examples, but keep it at my level.", quick_action="Explain & Examples")
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

# Chat input
if prompt := st.chat_input("Type your question here... (any language)", key="main_chat_input"):
    process_user_input(prompt)
    st.rerun()

# Temporarily disabled: Inject floating message history panel
# st.markdown(history_panel_html, unsafe_allow_html=True)

# Temporarily disabled: Inject JavaScript for click handlers
# scroll_js = f'''
# <script>
# ...
# </script>
# '''
# components.html(scroll_js, height=0, scrolling=False)